# -*- coding: utf-8 -*-
"""prophet-model-spec 재사용 docx 빌더 (데이터 주도형).

사용법:
    python docx_builder.py spec.json output.docx

spec.json 스키마 (모든 키 선택, 없으면 생략):
{
  "title":   "FIS Prophet 모델 명세서 (초안)",
  "product": "<상품명>",
  "tags":    "#변액연금 #VFA ...",
  "meta":    "영역 N개 · 변수 M개 · 근거: 약관/사업방법서(YYYY-MM-DD판)",
  "cover_note": "표지 하단 주의문구",
  "facts":   [ {"heading":"1.1 상품 개요", "headers":["항목","내용","근거"],
                "rows":[["...","...","..."]], "footnote":"※ ..."} , ... ],
  "areas":   [ {"title":"2. 영역 A — ...", "note":"영역 설명(선택)",
                "rows":[["변수명","타입","의미(근거)","포뮬라"], ...] }, ... ],
  "assumptions": {
     "sample":  [["1","변수/항목","영역","가정한 값","교체 근거"], ...],
     "approx":  [["1","항목","영역","근사 내용","확인 필요 원문"], ...],
     "missing": [["1","부재 항목","영향","비고"], ...]
  },
  "conventions": ["각주1", "각주2", ...],
  "warnings":    ["경고1", "경고2", ...],
  "footer": "생성: prophet-model-spec 스킬 · 근거: ..."
}

규칙:
- 포뮬라/값에 '// 샘플' 또는 '샘플'/'근사'가 있으면 자동으로 빨강 강조한다.
- assumptions.sample 표의 2번째 열(변수/항목)은 항상 샘플색으로 강조한다.
- 한 영역도 빠뜨리지 않는다(모델이 areas를 모두 채워 넣을 책임).
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KFONT = "맑은 고딕"
BLUE   = RGBColor(0x04, 0x3B, 0x72)
ORANGE = RGBColor(0xF5, 0x82, 0x20)
GREY   = RGBColor(0x55, 0x55, 0x55)
SAMPLE = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def _font(run, name=KFONT, size=9.5, bold=False, color=None, italic=False):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)


def _shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcpr.append(sh)


def _is_sample(v):
    s = str(v)
    return ("// 샘플" in s) or ("샘플" in s) or ("근사" in s)


class Builder:
    def __init__(self):
        self.doc = Document()
        st = self.doc.styles["Normal"]
        st.font.name = KFONT; st.font.size = Pt(9.5)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), KFONT)

    def h1(self, text, color=BLUE):
        p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        _font(p.add_run(text), size=15, bold=True, color=color)

    def h2(self, text, color=ORANGE):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
        _font(p.add_run(text), size=12, bold=True, color=color)

    def para(self, text, size=9.5, color=None, italic=False, bold=False):
        p = self.doc.add_paragraph()
        _font(p.add_run(text), size=size, color=color, italic=italic, bold=bold)

    def bullet(self, text, color=None):
        p = self.doc.add_paragraph(style="List Bullet")
        _font(p.add_run(text), size=9.5, color=color)

    def table(self, headers, rows, force_sample_col=None):
        t = self.doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hc = t.rows[0].cells
        for i, h in enumerate(headers):
            _shade(hc[i], "043B72")
            p = hc[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(p.add_run(str(h)), size=9, bold=True, color=WHITE)
        last = len(headers) - 1
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                p = cells[i].paragraphs[0]
                mono = (i == last and ("=" in str(val) or "//" in str(val)))
                col = SAMPLE if (_is_sample(val) or (force_sample_col is not None and i == force_sample_col)) else None
                _font(p.add_run(str(val)), name=("Consolas" if mono else KFONT), size=8.5, color=col)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)


def build(data, outpath):
    b = Builder()
    # 표지
    t = b.doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(t.add_run(data.get("title", "FIS Prophet 모델 명세서 (초안)")), size=20, bold=True, color=BLUE)
    if data.get("product"):
        p = b.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(data["product"]), size=14, bold=True, color=ORANGE)
    if data.get("tags"):
        p = b.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(data["tags"]), size=9, color=GREY)
    if data.get("meta"):
        p = b.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(data["meta"]), size=9.5, bold=True)
    if data.get("cover_note"):
        p = b.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(data["cover_note"]), size=8.5, color=GREY, italic=True)

    # 1. 추출된 상품 사실
    facts = data.get("facts", [])
    if facts:
        b.doc.add_paragraph()
        b.h1("1. 추출된 상품 사실 (약관 검증용)")
        b.para("사용자가 약관과 가장 먼저 대조해야 하는 부분. 조항번호는 약관/사업방법서 기준.",
               color=GREY, italic=True)
        for f in facts:
            if f.get("heading"):
                b.h2(f["heading"])
            b.table(f["headers"], f["rows"])
            if f.get("footnote"):
                b.para(f["footnote"], size=8.5, color=GREY)

    # 2~N. 영역별 변수
    areas = data.get("areas", [])
    if areas:
        b.page_break()
        for a in areas:
            b.h1(a["title"])
            if a.get("note"):
                b.para("▶ " + a["note"], size=9, color=GREY, italic=True)
            b.table(["변수명", "타입", "의미 (근거)", "포뮬라 (Prophet 유사)"], a["rows"])

    # 가정 사항 명세
    asm = data.get("assumptions")
    if asm:
        b.page_break()
        b.h1("가정 사항 명세 (Assumptions Log)", color=SAMPLE)
        b.para("약관에 근거하지 않고 '가정'한 모든 항목을 분류해 기록한다. 검증·교체 책임 추적용 핵심 섹션. "
               "분류: ⓐ 샘플 가정값 / ⓑ 구조 근사 / ⓒ 문서·데이터 부재.", color=GREY, italic=True)
        if asm.get("sample"):
            b.h2("ⓐ 샘플 가정값 — 산출방법서·경험자료로 교체 필수", color=SAMPLE)
            b.table(["#", "변수/항목", "영역", "가정한 값·내용", "교체 근거(필요 자료)"],
                    asm["sample"], force_sample_col=1)
        if asm.get("approx"):
            b.h2("ⓑ 구조 근사 — 산출방법서 부재·OCR 손상으로 로직 추정", color=SAMPLE)
            b.table(["#", "항목", "영역", "근사한 내용", "확인 필요 원문"], asm["approx"])
        if asm.get("missing"):
            b.h2("ⓒ 문서·데이터 부재", color=SAMPLE)
            b.table(["#", "부재 항목", "영향", "비고"], asm["missing"])

    # 컨벤션 + 경고
    if data.get("conventions") or data.get("warnings"):
        b.page_break()
    if data.get("conventions"):
        b.h1("포뮬라 컨벤션 각주")
        for c in data["conventions"]:
            b.bullet(c)
    if data.get("warnings"):
        b.h1("경고 (필독)", color=SAMPLE)
        for w in data["warnings"]:
            b.bullet(w)
    if data.get("footer"):
        p = b.doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
        _font(p.add_run(data["footer"]), size=8, color=GREY, italic=True)

    b.save(outpath)
    return outpath


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("usage: python docx_builder.py spec.json output.docx"); sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as fh:
        data = json.load(fh)
    print("SAVED:", build(data, sys.argv[2]))
